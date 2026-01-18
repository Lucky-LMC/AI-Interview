# AI智能面试辅助系统V1.0，作者刘梦畅
"""
Word 解析工具
提供 Word 文件的文本提取功能
"""
import os
from langchain_core.tools import tool

try:
    from docx import Document
except ImportError:
    Document = None


@tool
def word_parser_tool(file_path: str) -> str:
    """
    解析 Word 文件并提取文本内容
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    if Document is None:
        return "错误：请先安装 python-docx: pip install python-docx"
    
    if not os.path.exists(file_path):
        return f"错误：文件不存在 - {file_path}"
    
    try:
        doc = Document(file_path)
        
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        full_text = "\n".join(text_content)
        return full_text.strip() if full_text else "警告：Word 文件为空或无法提取文本"
        
    except Exception as e:
        return f"错误：Word 解析失败 - {str(e)}"
