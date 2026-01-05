# AI模拟面试系统v1.0，作者刘梦畅
"""
简历解析节点
使用 Agent + 工具自动解析文档（PDF/Word）
"""
import os
from langchain_core.messages import AIMessage
from langchain_core.tools import tool
from backend.models.state import InterviewState
from .llm_helper import get_shared_llm
from langgraph.prebuilt import create_react_agent
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from docx import Document
except ImportError:
    Document = None


# ========== 文档解析工具 ==========

@tool
def pdf_parser_tool(file_path: str) -> str:
    """
    解析 PDF 文件并提取文本内容
    
    从 PDF 文件路径中提取所有页面的文本内容。
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容，如果失败返回错误信息
    """
    if PyPDF2 is None:
        return "错误：请先安装 PyPDF2: pip install PyPDF2"
    
    if not file_path:
        return "错误：未提供文件路径"
    
    if not os.path.exists(file_path):
        return f"错误：文件不存在 - {file_path}"
    
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            # 提取所有页面的文本
            text_content = []
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_content.append(text)
            
            # 合并所有文本
            full_text = "\n".join(text_content)
            
            if full_text:
                return full_text.strip()
            else:
                return "警告：PDF 文件为空或无法提取文本"
            
    except Exception as e:
        error_msg = f"PDF 解析错误: {str(e)}"
        print(error_msg)
        return f"错误：PDF 解析失败 - {str(e)}"


@tool
def word_parser_tool(file_path: str) -> str:
    """
    解析 Word 文件并提取文本内容
    
    Args:
        file_path: 文件路径
        
    Returns:
        str: 提取的文本内容
    """
    if Document is None:
        return "错误：请先安装 python-docx: pip install python-docx"
    
    if not os.path.exists(file_path):
        return f"错误：文件不存在 - {file_path}"
    
    try:
        doc = Document(file_path)
        
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # 提取表格内容
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        full_text = "\n".join(text_content)
        return full_text.strip() if full_text else "警告：Word 文件为空或无法提取文本"
        
    except Exception as e:
        return f"错误：Word 解析失败 - {str(e)}"


# Agent 系统提示词
DOCUMENT_AGENT_PROMPT = """
你是一个专业的文档解析助手。

1. 收到文件路径后，先根据文件扩展名自行判断类型：
   - .pdf → 使用 pdf_parser_tool({"file_path": 路径})
   - .docx/.doc → 使用 word_parser_tool({"file_path": 路径})
2. 如果扩展名不是上述格式，直接说明暂不支持
3. 解析完成后只返回提取的纯文本内容
"""


def create_document_agent():
    """创建文档解析 Agent"""
    if create_react_agent is None:
        raise RuntimeError("请先安装 langgraph: pip install langgraph")
    
    tools = [pdf_parser_tool, word_parser_tool]
    llm = get_shared_llm()
    
    # 使用 langgraph 的 create_react_agent
    return create_react_agent(
        model=llm,
        tools=tools,
        prompt=DOCUMENT_AGENT_PROMPT
    )

document_agent = create_document_agent()

def parse_resume_node(state: InterviewState) -> InterviewState:
    """
    简历解析节点：使用 Agent + 工具自动解析文档（PDF/Word）
    
    Agent 会根据文件扩展名自动选择合适的工具进行解析
    """
    # 如果 resume_text 已经存在，直接返回
    if state.get('resume_text'):
        return state
    
    resume_path = state.get('resume_path', '')
    
    if not resume_path:
        print("[parse_resume_node] 警告: 没有提供简历文件路径")
        return state
    
    if not os.path.exists(resume_path):
        print(f"[parse_resume_node] 警告: 简历文件不存在: {resume_path}")
        return state
        
    # 使用 Agent 解析文档
    try:
        print(f"[parse_resume_node] 开始解析: {resume_path}")
        
        # 调用 Agent 解析
        user_message = f"请解析这个文件：{resume_path}"
        inputs = {
            "messages": [
                {"role": "user", "content": user_message}
            ]
        }
        
        result = document_agent.invoke(inputs)
        
        # 从 Agent 响应中提取文本内容
        resume_text = ""
        if result and "messages" in result:
            messages = result["messages"]
            for msg in reversed(messages):
                if isinstance(msg, AIMessage):
                    content = getattr(msg, "content", "")
                    if isinstance(content, str) and content.strip():
                        resume_text = content.strip()
                        break
        
        # 检查解析结果
        if not resume_text or resume_text.startswith("解析失败") or resume_text.startswith("错误"):
            print(f"[parse_resume_node] Agent 返回无效结果: {resume_text}")
            return state
        
        print(f"[parse_resume_node] 解析成功，文本长度: {len(resume_text)}")
        
        # 更新状态
        new_state = state.copy()
        new_state['resume_text'] = resume_text
        return new_state
            
    except Exception as e:
        print(f"[parse_resume_node] 解析简历失败: {e}", exc_info=True)
        raise
